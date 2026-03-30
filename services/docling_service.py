"""
Docling-backed parsing layer for KB file ingestion.

Sprint 1 scope:
- define a stable parsing interface for document ingestion
- support parser modes: legacy_only / docling_with_legacy_fallback / docling_only
- keep existing bot behaviour safe by falling back to the legacy parser when
  rollout policy allows it

The returned object is intentionally simple for now: normalized text plus
parser metadata. Later sprints will extend it with structure-aware fields.
"""

from __future__ import annotations

from dataclasses import asdict
from importlib.metadata import version as pkg_version
from dataclasses import dataclass, field
from io import BytesIO
import json
import os
import subprocess
import sys
import tempfile
from typing import Any

from services import kb_rollout

DOCLING_SUPPORTED_EXTS = {
    "pdf",
    "docx",
    "pptx",
    "html",
    "htm",
    "md",
    "csv",
    "xlsx",
    "xml",
    "latex",
    "tex",
    "asciidoc",
    "adoc",
    "vtt",
}
DOCLING_PDF_SUBPROCESS_TIMEOUT_SEC = 180


@dataclass
class ParsedDocument:
    filename: str
    text: str
    parser_backend: str
    parser_mode: str
    parser_version: str | None = None
    source_format: str | None = None
    page_count: int | None = None
    has_tables: bool = False
    has_headings: bool = False
    structure: dict[str, Any] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)
    fallback_used: bool = False
    fallback_reason: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


def is_docling_supported_filename(filename: str) -> bool:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return ext in DOCLING_SUPPORTED_EXTS


def _extract_markdown_headings(text: str) -> list[str]:
    headings: list[str] = []
    for line in (text or "").splitlines():
        stripped = line.strip()
        if not stripped.startswith("#"):
            continue
        heading = stripped.lstrip("#").strip()
        if heading:
            headings.append(heading)
    return headings


def _guess_has_tables(text: str) -> bool:
    for line in (text or "").splitlines():
        stripped = line.strip()
        if stripped.count("|") >= 2:
            return True
    return False


def _create_docling_converter(filename: str):
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.document_converter import DocumentConverter, PdfFormatOption

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        # For most KB uploads we receive digital PDFs with embedded text.
        # Disabling OCR avoids heavy model downloads and multi-minute CPU stalls.
        pdf_options = PdfPipelineOptions(do_ocr=False)
        return DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_options),
            }
        )
    return DocumentConverter()


def _get_docling_pdf_timeout_sec(data_size_bytes: int | None = None) -> int:
    raw = os.getenv("KB_DOCLING_PDF_TIMEOUT_SEC", str(DOCLING_PDF_SUBPROCESS_TIMEOUT_SEC))
    try:
        base_timeout = max(30, int(raw))
    except (TypeError, ValueError):
        base_timeout = DOCLING_PDF_SUBPROCESS_TIMEOUT_SEC

    if not data_size_bytes or data_size_bytes <= 0:
        return base_timeout

    size_mb = data_size_bytes / (1024 * 1024)
    adaptive_timeout = base_timeout + int(size_mb * 20)
    return min(420, max(base_timeout, adaptive_timeout))


def _project_root() -> str:
    return os.path.dirname(os.path.dirname(__file__))


def _parse_with_docling_in_subprocess(filename: str, data: bytes, *, parser_mode: str) -> ParsedDocument:
    with tempfile.TemporaryDirectory(prefix="blabber-docling-") as tmp_dir:
        input_path = os.path.join(tmp_dir, "input.bin")
        output_path = os.path.join(tmp_dir, "output.json")
        with open(input_path, "wb") as f:
            f.write(data)

        env = os.environ.copy()
        project_root = _project_root()
        existing_pythonpath = env.get("PYTHONPATH", "")
        env["PYTHONPATH"] = (
            project_root if not existing_pythonpath else project_root + os.pathsep + existing_pythonpath
        )
        cmd = [
            sys.executable,
            "-m",
            "services.docling_service",
            "--worker",
            input_path,
            output_path,
            filename,
            parser_mode,
        ]
        try:
            completed = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                env=env,
                timeout=_get_docling_pdf_timeout_sec(len(data)),
                cwd=project_root,
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError(f"Docling PDF worker timeout after {int(exc.timeout)}s")

        if completed.returncode != 0:
            stderr = (completed.stderr or completed.stdout or "").strip()
            raise RuntimeError(
                "Docling PDF worker failed"
                + (f": {stderr[:300]}" if stderr else "")
            )

        if not os.path.exists(output_path):
            raise RuntimeError("Docling PDF worker did not produce output")

        with open(output_path, "r", encoding="utf-8") as f:
            payload = json.load(f)
        return ParsedDocument(**payload)


def parse_document(
    filename: str,
    data: bytes,
    *,
    parser_mode: str | None = None,
) -> ParsedDocument:
    """
    Parse one document into normalized text using the configured parser mode.
    """
    mode = (parser_mode or kb_rollout.get_doc_parser_mode()).strip().lower()
    if mode not in {"legacy_only", "docling_with_legacy_fallback", "docling_only"}:
        mode = "legacy_only"

    if mode == "legacy_only":
        return _parse_with_legacy(filename, data, parser_mode=mode)

    if not is_docling_supported_filename(filename):
        if mode == "docling_only":
            ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "unknown"
            raise ValueError(
                f"Docling parser пока не поддерживает формат .{ext} в этом пайплайне."
            )
        parsed = _parse_with_legacy(filename, data, parser_mode=mode)
        parsed.fallback_used = True
        parsed.fallback_reason = "docling_format_not_supported"
        parsed.warnings.append("Docling fallback: unsupported format for Docling")
        parsed.metadata["docling_unsupported_format"] = True
        return parsed

    if mode == "docling_only":
        return _parse_with_docling(filename, data, parser_mode=mode)

    try:
        return _parse_with_docling(filename, data, parser_mode=mode)
    except Exception as exc:
        if not kb_rollout.should_continue_after_docling_failure():
            raise
        parsed = _parse_with_legacy(filename, data, parser_mode=mode)
        parsed.fallback_used = True
        parsed.fallback_reason = str(exc)
        parsed.warnings.append(f"Docling fallback: {exc}")
        parsed.metadata["docling_error"] = str(exc)
        return parsed


def _parse_with_docling(filename: str, data: bytes, *, parser_mode: str) -> ParsedDocument:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return _parse_with_docling_in_subprocess(filename, data, parser_mode=parser_mode)
    return _parse_with_docling_inprocess(filename, data, parser_mode=parser_mode)


def _parse_with_docling_inprocess(filename: str, data: bytes, *, parser_mode: str) -> ParsedDocument:
    try:
        from docling.datamodel.base_models import DocumentStream
    except ImportError:
        raise ValueError(
            "Для Docling parsing нужен пакет <code>docling</code>.\n"
            "Установи: <code>pip install docling</code>"
        )

    source = DocumentStream(name=filename, stream=BytesIO(data))
    converter = _create_docling_converter(filename)
    result = converter.convert(source)
    doc = getattr(result, "document", None)
    if doc is None:
        raise ValueError("Docling не вернул document result")

    text = ""
    if hasattr(doc, "export_to_markdown"):
        text = (doc.export_to_markdown() or "").strip()
    if not text and hasattr(doc, "export_to_text"):
        text = (doc.export_to_text() or "").strip()
    if not text:
        raise ValueError("Docling не смог извлечь текст из документа")

    headings = _extract_markdown_headings(text)
    has_tables = _guess_has_tables(text)

    status = getattr(result, "status", None)
    status_name = getattr(status, "name", str(status) if status is not None else "")
    input_obj = getattr(result, "input", None)
    input_fmt = getattr(input_obj, "format", None)
    source_format = getattr(input_fmt, "name", str(input_fmt) if input_fmt is not None else None)
    page_count = None
    pages = getattr(doc, "pages", None)
    try:
        if pages is not None:
            page_count = len(pages)
    except Exception:
        page_count = None

    warnings: list[str] = []
    errors = getattr(result, "errors", None) or []
    if errors:
        warnings.extend(str(err) for err in errors[:5])
    if status_name and status_name not in {"SUCCESS", "ConversionStatus.SUCCESS"}:
        warnings.append(f"Docling status: {status_name}")

    return ParsedDocument(
        filename=filename,
        text=text,
        parser_backend="docling",
        parser_mode=parser_mode,
        parser_version=pkg_version("docling"),
        source_format=source_format,
        page_count=page_count,
        has_tables=has_tables,
        has_headings=bool(headings),
        structure={"headings": headings[:50]},
        warnings=warnings,
        metadata={
            "docling_status": status_name,
            "ocr_enabled": False if filename.lower().endswith(".pdf") else None,
        },
    )


def _docling_worker_main(argv: list[str]) -> int:
    if len(argv) != 5 or argv[0] != "--worker":
        print(
            "Usage: python -m services.docling_service --worker <input> <output> <filename> <parser_mode>",
            file=sys.stderr,
        )
        return 2

    _flag, input_path, output_path, filename, parser_mode = argv
    try:
        with open(input_path, "rb") as f:
            data = f.read()
        parsed = _parse_with_docling_inprocess(filename, data, parser_mode=parser_mode)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(asdict(parsed), f, ensure_ascii=True)
        return 0
    except Exception as exc:
        print(str(exc), file=sys.stderr)
        return 1


def _parse_with_legacy(filename: str, data: bytes, *, parser_mode: str) -> ParsedDocument:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in ("txt", "md"):
        for enc in ("utf-8", "cp1251", "latin-1"):
            try:
                return ParsedDocument(
                    filename=filename,
                    text=data.decode(enc),
                    parser_backend="legacy",
                    parser_mode=parser_mode,
                    parser_version=None,
                    source_format=ext or None,
                )
            except UnicodeDecodeError:
                continue
        return ParsedDocument(
            filename=filename,
            text=data.decode("utf-8", errors="replace"),
            parser_backend="legacy",
            parser_mode=parser_mode,
            parser_version=None,
            source_format=ext or None,
        )

    if ext == "pdf":
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(BytesIO(data))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n\n".join(p for p in pages if p.strip())
            return ParsedDocument(
                filename=filename,
                text=text,
                parser_backend="legacy",
                parser_mode=parser_mode,
                parser_version=None,
                source_format=ext,
            )
        except ImportError:
            raise ValueError(
                "Для PDF нужен пакет pypdf.\n"
                "Установи: <code>pip install pypdf</code>\n"
                "Или загрузи TXT-версию документа."
            )
        except Exception as exc:
            raise ValueError(f"Не удалось прочитать PDF: {exc}")

    if ext in ("docx", "doc"):
        try:
            import docx as docx_lib  # type: ignore

            doc = docx_lib.Document(BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return ParsedDocument(
                filename=filename,
                text=text,
                parser_backend="legacy",
                parser_mode=parser_mode,
                parser_version=None,
                source_format=ext,
            )
        except ImportError:
            raise ValueError(
                "Для DOCX нужен пакет python-docx.\n"
                "Установи: <code>pip install python-docx</code>\n"
                "Или загрузи TXT-версию документа."
            )
        except Exception as exc:
            raise ValueError(f"Не удалось прочитать DOCX: {exc}")

    return ParsedDocument(
        filename=filename,
        text=data.decode("utf-8", errors="replace"),
        parser_backend="legacy",
        parser_mode=parser_mode,
        parser_version=None,
        source_format=ext or None,
    )


if __name__ == "__main__":
    raise SystemExit(_docling_worker_main(sys.argv[1:]))
